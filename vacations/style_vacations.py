


from tkinter import *
from tkinter import ttk
from PIL import Image,ImageTk
from tkinter import messagebox
from tkinter import filedialog
from threading import Thread
from tkcalendar import Calendar, DateEntry
from datetime import datetime
from babel.dates import format_date, parse_date, get_day_names, get_month_names
from babel.numbers import * 
from vacations.cont_database_vacations import DataBaseVacations
import os 
import csv
import docx
import sqlite3

class StyleVacations:
	def __init__(self,master):
		self.master = Toplevel()
		self.master.geometry("1180x800+300+100")
		self.master.resizable(width=False,height=False)
		self.master.configure(background='#EAF2F8')
		self.master.title("نظام الإجازات")


	
		self.database_vavations = os.path.abspath(".")
		self.joins_database_vacations = os.path.join(self.database_vavations,"database_vavations.db")


		self.menubar = Menu(self.master)
		self.file = Menu(self.menubar,tearoff=False)
		self.file.add_command(label='Save As')

		self.menubar.add_cascade(label='حفظ البيانات',command=self.save_vacations_employees)
		self.master.config(menu=self.menubar)


		self.class_database_vacations = DataBaseVacations(self.master)


		
		self.title_label = Label(self.master,

			text="الإجازات" , width=71,font=("Libre Baskerville, serif;",20),
			background="#E8E4FA")

		self.title_label.pack(fill=X)

		self.frame_info = Frame(self.master ,background='#E7F9F4',
		relief=RIDGE , bd=2 , width=1155,height=272).place(x=20,y=123)




		self.run_add_frame_search = Thread(target=self.add_frame_search,args=(),)
		self.run_add_frame_search.start()


		self.run_add_button_search_employee = Thread(target=self.add_button_search_employee,args=(),)
		self.run_add_button_search_employee.start()

		self.run_add_email = Thread(target=self.add_email,args=(),)
		self.run_add_email.start()


		self.run_add_name = Thread(target=self.add_name,args=(),)
		self.run_add_name.start()


		self.run_add_id = Thread(target=self.add_id,args=(),)
		self.run_add_id.start()


		self.run_add_section = Thread(target=self.add_section,args=(),)
		self.run_add_section.start()



		self.run_add_type_vacation = Thread(target=self.add_type_vacation,args=(),)
		self.run_add_type_vacation.start()



		self.run_add_number_vacation = Thread(target=self.add_number_vacation,args=(),)
		self.run_add_number_vacation.start()



		self.run_add_date_vacation = Thread(target=self.add_date_vacation,args=(),)
		self.run_add_date_vacation.start()



		

		self.run_add_coincidence_day = Thread(target=self.add_coincidence_day,args=(),)
		self.run_add_coincidence_day.start()



		
		self.run_add_frame_buttons = Thread(target=self.add_frame_buttons,args=(),)
		self.run_add_frame_buttons.start()




		# =================== Run Funcatins Buttons ================================

		self.run_function_button_insert = Thread(target=self.function_insert,args=(),)
		self.run_function_button_insert.start()



		self.run_function_button_update = Thread(target=self.function_update,args=(),)
		self.run_function_button_update.start()




		self.run_function_button_delete = Thread(target=self.function_delete,args=(),)
		self.run_function_button_delete.start()



		
		self.run_function_button_view = Thread(target=self.function_view,args=(),)
		self.run_function_button_view.start()




		self.run_function_button_empty = Thread(target=self.function_empty,args=(),)
		self.run_function_button_empty.start()





	def add_frame_search(self):
		""" The Function Add Entry Search Employee """

		self.frame_search = Frame(self.master,

			background='#d9ddf7', relief=RIDGE , bd=2 , width=1155,height=80
			).place(x=20,y=43)

		
		self.lable_search = Label(self.master,text="الرجاء إدخال رقم الموظف للبحث عنه",
			font=("bold",20),
			background='#d9ddf7' , 
			width=23
			).place(x=800,y=58)



		self.entry_search = Entry(
			self.master , width=50 ,font=('Bold Oblique',16) , relief=RAISED , bd=1 ,justify='center',
			background='white',textvariable=self.class_database_vacations.variable_search_employee_vacations)
		self.entry_search.place(x=200,y=62)








	def add_button_search_employee(self):
		self.image_search = Image.open('vacations\\images\\search.png')
		self.image_search = self.image_search.resize((35,35))
		self.insert_image_search = ImageTk.PhotoImage(self.image_search)

		self.button_search = Button(
		self.master , background='#F7F7F9' , width=90 , height=30 , text='بــحث',font=('Bold Oblique',18) ,
		image=self.insert_image_search , compound='left',padx=20,relief=FLAT , bd=1,command=self.class_database_vacations.search_employee_vacations
		).place(x=50,y=57)



	
		


	def add_email(self):
		self.label_email = Label(self.master,text='البريد الإلكتروني',font=("bold",20),bg='#E7F9F4')
		self.label_email.place(x=1006,y=137)
		self.email = Entry(self.master , width=27,font=('Bold Oblique',18),
			relief=RIDGE , bd=1,justify='center',textvariable=self.class_database_vacations.var_email).place(x=638,y=142)


	def add_name(self):
		self.label_name = Label(self.master,text='إسم الموظف',font=("bold",20),bg='#E7F9F4')
		self.label_name.place(x=480,y=140)
		self.name = Entry(self.master , width=31,font=('Bold Oblique',18),
			relief=RIDGE , bd=1,justify='center',textvariable=self.class_database_vacations.var_name).place(x=49,y=142)



	def add_id(self):
		self.label_id = Label(self.master,text='الرقم الوظيفي',font=("bold",20),bg='#E7F9F4')
		self.label_id.place(x=1010,y=198)
		self.id = Entry(self.master , width=27,font=('Bold Oblique',18),
			relief=RIDGE , bd=1,justify='center',textvariable=self.class_database_vacations.var_id).place(x=638,y=203)


	

	def add_section(self):
		self.label_section = Label(self.master,text='القسم الوظيفي',font=("bold",20),bg='#E7F9F4')
		self.label_section.place(x=478,y=200)
		self.section = Entry(self.master , width=31,font=('Bold Oblique',18),
			relief=RIDGE , bd=1,justify='center',textvariable=self.class_database_vacations.var_section).place(x=49,y=200)



	def add_type_vacation(self):
		self.label_vacation = Label(self.master,text='نوع الإجازة',font=("bold",20),bg='#E7F9F4')
		self.label_vacation.place(x=1010,y=270)

		self.list_type_vacation = ("إجازة دورية" ,"إجازة عارضة" , "إجازة مرضية" , "إجازة وضع" , "إجازة رعاية الأبناء")

		self.vacation = ttk.Combobox(self.master,
			width=26 , font=('Bold Oblique',18) ,justify='center',values=self.list_type_vacation,
			textvariable=self.class_database_vacations.var_type_coincidence)
		self.vacation.place(x=638,y=270)

		self.vacation.option_add('*TCombobox*Listbox.Justify', 'center') 




	def add_number_vacation(self):
		self.label_number_vacation = Label(self.master,text='عدد أيام الإجازة',font=("bold",20),bg='#E7F9F4')
		self.label_number_vacation.place(x=478,y=270)

		self.number_vacation = Entry(self.master , width=31,font=('Bold Oblique',18),
			relief=RIDGE , bd=1,justify='center',textvariable=self.class_database_vacations.var_number_days_vacation).place(x=49,y=270)





	
	
	
	

	def add_date_vacation(self):
		self.label_date_vacation = Label(self.master,text='تأريخ الإجازة',font=("bold",20),bg='#E7F9F4')
		self.label_date_vacation.place(x=1010,y=343)


		self.date_vacation = DateEntry(self.master,selectmode='day',year=datetime.now().year,
			month=datetime.now().month,day=datetime.now().day,
			font=('Bold Oblique',18),
			justify='center',date_pattern="yyyy-mm-dd", foreground="black",width=26,
			headersforeground="black", selectforeground="black",textvariable=self.class_database_vacations.var_date_coincidence)

		self.date_vacation.place(x=638,y=343)


	


	def add_coincidence_day(self):
		self.label_coincidence_day = Label(self.master,text='المصادف يوم',font=("bold",20),bg='#E7F9F4')
		self.label_coincidence_day.place(x=478,y=341)

		self.day = Entry(self.master , width=31,font=('Bold Oblique',18),
			relief=RIDGE , bd=1,justify='center',textvariable=self.class_database_vacations.var_coincidence_day).place(x=49,y=341)





	def add_frame_buttons(self):
		""" Add Frame Buttons Controls Data """
		self.frame_buttons = Frame(self.master,

			background='#FCEEE6', relief=RIDGE , bd=2 , width=1155,height=80
			).place(x=20,y=395)



	def function_insert(self):
		""" Insert Data In DataBase """
		self.image_insert = Image.open('vacations\\images\\add_user.png')
		self.image_insert = self.image_insert.resize((41,41))
		self.image_save = ImageTk.PhotoImage(self.image_insert)

		

		self.button_insert = Button(
			self.master , background='white' , width=145 , height=40 , text='حفظ البيانات',font=('Bold Oblique',18),
			relief=FLAT , bd=1 ,image=self.image_save , compound='left',padx=20,command=self.class_database_vacations.check_data
			).place(x=970,y=410)



	def function_update(self):
		""" Update Data In DataBase """
		self.image_update = Image.open('vacations\\images\\update_user.png')
		self.image_update = self.image_update.resize((41,41))
		self.image_save_update = ImageTk.PhotoImage(self.image_update)

		

		self.button_update = Button(
			self.master , background='white' , width=145 , height=40 , text='تحديث البيانات',font=('Bold Oblique',18),
			relief=FLAT , bd=1 ,image=self.image_save_update , compound='left',padx=20,command=self.class_database_vacations.update_data_vacations
			).place(x=730,y=410)




	def function_delete(self):
		""" Delete Data In DataBase """
		self.image_delete = Image.open('vacations\\images\\delete_user.png')
		self.image_delete = self.image_delete.resize((41,41))
		self.image_save_delete = ImageTk.PhotoImage(self.image_delete)

		

		self.button_delete = Button(
			self.master , background='white' , width=145 , height=40 , text='حذف البيانات',font=('Bold Oblique',18),
			relief=FLAT , bd=1 ,image=self.image_save_delete , compound='left',padx=20,command=self.class_database_vacations.delete_data_vacations
			).place(x=500,y=410)



	def function_view(self):
		""" view Data In DataBase """
		self.image_view = Image.open('vacations\\images\\view_data.png')
		self.image_view = self.image_view.resize((39,39))
		self.image_save_view = ImageTk.PhotoImage(self.image_view)

		

		self.button_view = Button(
			self.master , background='white' , width=145 , height=40 , text='عرض البيانات',font=('Bold Oblique',18),
			relief=FLAT , bd=1 ,image=self.image_save_view , compound='left',padx=20,command=self.class_database_vacations.show_data_vacations
			).place(x=270,y=410)




	def function_empty(self):
		""" view Data In DataBase """
		self.image_empty = Image.open('vacations\\images\\empty_data.png')
		self.image_empty = self.image_empty.resize((39,39))
		self.image_save_empty = ImageTk.PhotoImage(self.image_empty)

		

		self.button_empty = Button(
			self.master , background='white' , width=145 , height=40 , text='إفراغ الحقول',font=('Bold Oblique',18),
			relief=FLAT , bd=1 ,image=self.image_save_empty , compound='left',padx=20,command=self.class_database_vacations.clean_data_vacations
			).place(x=50,y=410)






	def save_vacations_employees(self):
		self.title = "قائمة الإجازات بتأريخ"
		self.date_time = str(datetime.now().date())


		self.doc = docx.Document()
		self.doc.add_heading(f"{self.title} - {self.date_time}")
		

		self.table = self.doc.add_table(rows=1, cols=9)
		self.row = self.table.rows[0].cells
		self.table.style = "Colorful List Accent 1"
		#self.table.style = 'Colorful List'

		self.table_header = [
							"المصادف يوم"  , "عدد أيام الإجازة" , "القسم الوظيفي" 
							, "إسم الموظف" , "تأريخ الإجازة" , "نوع الإجازة" , "الرقم الوظيفي" , "البريد الإلكتروني"

							]



		for i in range(8):
			self.table.rows[0].cells[i].text = self.table_header[i]


		self.con = sqlite3.connect("vacations\\database_vacations.db")
		self.curs = self.con.cursor()

		self.res = self.curs.execute("SELECT coincidence_day ,number_days_vacation , Section, Name,date_coincidence,type_coincidence , ID ,Email FROM VacationsEmployee ")
		self.data_sql = self.res.fetchall()

		for (
			self.coincidence_day , self.number_days_vacation , self.section ,
			self.name , self.date_coincidence , self.type_coincidence ,
			self.id , self.eamil
			) in self.data_sql:

			self.row = self.table.add_row().cells
			self.row[0].text = self.coincidence_day
			self.row[1].text = self.number_days_vacation
			self.row[2].text = self.section
			self.row[3].text = self.name
			self.row[4].text = str(self.date_coincidence)
			self.row[5].text = self.type_coincidence

			self.row[6].text = str(self.id)
			self.row[7].text = self.eamil
		



		try :
		
			self.path = filedialog.asksaveasfilename(initialdir = "C:\\",title = "Save As",filetypes = ("docx files",'docx files(*.*)'))
			self.doc.save(self.path)
			self.con.commit()
		except :
			pass