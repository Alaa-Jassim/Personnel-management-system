




from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from tkcalendar import Calendar, DateEntry
from PIL import Image , ImageTk
from datetime import datetime
import os , sys
from list_employee.database_list_employee import DataBaseListEmployee
from list_employee.treeview_list_employee import TreeviewListEmployee
import sqlite3
from tkinter import filedialog
import csv
import docx
from babel.dates import format_date, parse_date, get_day_names, get_month_names
from babel.numbers import * 

class StyleListEmployee:
	def __init__(self,master):
		self.master = Toplevel()
		self.master.geometry("1180x600+300+100")
		self.master.resizable(width=False,height=False)
		self.master.configure(background='#EAF2F8')
		self.master.title("أقسام الموظفين")
		self.master.iconbitmap('list_employee\\images\\sections_em.ico')



		self.menubar = Menu(self.master)
		self.file = Menu(self.menubar,tearoff=False)
		self.file.add_command(label='Save As')

		self.menubar.add_cascade(label='حفظ البيانات',command=self.open_file)
		self.master.config(menu=self.menubar)








		self.class_database = DataBaseListEmployee(self.master)
		self.class_treeview = TreeviewListEmployee(self.master)

	

		self.lable_title = Label(self.master,background='#E8E4FA',
								text="أقسام الموظفين",
								width=71,
								font=("Libre Baskerville, serif;",20))

		self.lable_title.place(x=20,y=4)



		self.frame_search_section = Frame(self.master ,background='#EBF5FB',
		relief=RIDGE , bd=2 , width=1143,height=90).place(x=20,y=43)

		self.add_label_id()
		self.add_label_section()
		self.add_label_name()
		self.add_buttin_show()







	def add_label_id(self):
		self.label_id = Label(self.master ,
		 					text="الرقم الوظيفي",
							font=("bold",20),
							background='#EBF5FB' ,
							width=8 )
		self.label_id.place(x=1020,y=63)
		self.entry_id = Entry(self.master , width=16 , font=('Bold Oblique',14),
			relief=RIDGE , bd=1,justify='center',textvariable=self.class_database.var_id_list).place(x=834,y=71)





	def add_label_section(self):
		self.label_section = Label(self.master ,
		 					text="القسم",
							font=("bold",20),
							background='#EBF5FB' ,
							width=8 )
		self.label_section.place(x=775,y=65)


		self.entry_sections = Entry(self.master , width=20 , font=('Bold Oblique',14),relief=RIDGE ,
		 bd=1,justify='center',textvariable=self.class_database.var_section_list).place(x=575,y=71)
		

	





	def add_label_name(self):
		self.label_name = Label(self.master ,
			 					text="إسم الموظف",
								font=("bold",20),
								background='#EBF5FB' ,
								width=8 )
		self.label_name.place(x=420,y=65)
		self.entry_name = Entry(self.master , width=20 , font=('Bold Oblique',14),
			relief=RIDGE , bd=1,justify='center',textvariable=self.class_database.var_name_list).place(x=185,y=71)





	def add_buttin_show(self):
		self.image_search = Image.open('folder_attendance_departure\\images\\search.png')
		self.image_search = self.image_search.resize((35,35))
		self.insert_image_search = ImageTk.PhotoImage(self.image_search)

		self.button_show = Button(self.master , 
								text="عرض" , 
								background='#F7F7F9',
								width=80 , 
								height=30 ,
								font=('Bold Oblique',18),
								compound='left',padx=20,relief=FLAT , bd=1,
								image=self.insert_image_search,command=self.class_database.get_data_list
								)

		self.button_show.place(x=35,y=65)





	def open_file(self):
		self.path_aps = os.path.abspath(".") 
		self.file_database = os.path.join(self.path_aps, "DataBaseEmployees.db")

		self.title = "قائمة أقسام الموظفين بتأريخ"
		self.date_time = str(datetime.now().date())


		self.doc = docx.Document()
		self.doc.add_heading(f"{self.title} - {self.date_time}")
		self.table_header = ["الإسم"   , "القسم"   , "الرقم الوظيفي"]

		self.table = self.doc.add_table(rows=1, cols=3)
		self.row = self.table.rows[0].cells
		self.table.style = "Colorful List Accent 4"
		


		for i in range(3):
			self.table.rows[0].cells[i].text = self.table_header[i]


		self.con = sqlite3.connect(self.file_database)
		self.curs = self.con.cursor()

		self.res = self.curs.execute('SELECT Name , Section , ID FROM Employees')
		self.data_sql = self.res.fetchall()

		for self.name , self.section , self.id in self.data_sql:

			self.row = self.table.add_row().cells
			self.row[0].text = self.name
			self.row[1].text = self.section
			self.row[2].text = str(self.id)


		try :
		
			self.path = filedialog.asksaveasfilename(initialdir = ".",title = "Save As",filetypes = ("docx files",'docx files(*.*)'))
			self.doc.save(self.path)
			self.con.commit()
		except :
			pass
	  

  
