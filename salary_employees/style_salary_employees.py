

from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from tkcalendar import Calendar, DateEntry
from PIL import Image , ImageTk
from datetime import datetime
import os , sys
import docx
from datetime import datetime
import os , sys , shutil
import sqlite3
from threading import Thread
from tkinter import filedialog
from salary_employees.database_salary import DatabaseSalaryEmployees
from babel.dates import format_date, parse_date, get_day_names, get_month_names
from babel.numbers import * 

class StyleSalaryEmployee:
	def __init__(self,master):
		self.master = Toplevel()
		self.master.geometry("1180x600+300+100")
		self.master.resizable(width=False,height=False)
		self.master.configure(background='#EAF2F8')
		self.master.title("رواتب الموظفين")
		self.master.iconbitmap('salary_employees\\images\\icon_app.ico')


		self.menubar = Menu(self.master)
		self.file = Menu(self.menubar,tearoff=False)
		self.file.add_command(label='Save As')

		self.menubar.add_cascade(label='حفظ البيانات',command=self.open_file)
		self.master.config(menu=self.menubar)


		self.lable_title = Label(self.master,background='#E8E4FA',
								text="رواتب الموظفين",
								width=71,
								font=("Libre Baskerville, serif;",20))

		self.lable_title.place(x=20,y=4)





		self.class_database_salary = DatabaseSalaryEmployees(self.master)




		self.frame_search_salary = Frame(self.master ,background='#d9ddf7',
		relief=RIDGE , bd=2 , width=1143,height=90).place(x=20,y=43)

		self.add_label_search_salary()
		self.add_button_salary()




	def add_label_search_salary(self):

		self.label_salary = Label(self.master,
			text="إدخل الرقم الوظيفي أو البريد الإلكتروني للبحث" ,
			width=28,
			font=("Libre Baskerville, serif;",19),background='#d9ddf7')

		self.label_salary.place(x=734,y=69)


		self.entry_name_or_id = Entry(self.master , width=46 , font=('Bold Oblique',14),relief=RIDGE ,
		 bd=1,justify='center',textvariable=self.class_database_salary.var_email_or_id_salary).place(x=220,y=71)
		




	def add_button_salary(self):
		self.image_salary = Image.open('salary_employees\\images\\salary_icon.png')
		self.image_salary = self.image_salary.resize((35,35))
		self.insert_image_salary = ImageTk.PhotoImage(self.image_salary)

		self.button_show = Button(self.master , 
								text="عرض" , 
								background='#F7F7F9',
								width=100 , 
								height=33 ,
								font=('Bold Oblique',18),
								compound='left',padx=20,relief=FLAT , bd=1,
								image=self.insert_image_salary,command=self.class_database_salary.get_data_salary
								)

		self.button_show.place(x=40,y=63)


	



	def open_file(self):
		self.title = "قائمة رواتب الموظفين بتأريخ  "
		self.date_time = str(datetime.now().date())


		self.doc = docx.Document()
		self.doc.add_heading(f"{self.title} - {self.date_time}")
		self.table_header = ["إسم الموظف" , "البريد الإلكتروني" , "القسم" , "الإجرة الشهرية" , "الرقم الوظيفي"]

		self.table = self.doc.add_table(rows=1, cols=5)
		self.row = self.table.rows[0].cells
		self.table.style = "Colorful List Accent 4"
		


		for i in range(5):
			self.table.rows[0].cells[i].text = self.table_header[i]


		self.con = sqlite3.connect('DataBaseEmployees.db')
		self.curs = self.con.cursor()

		self.res = self.curs.execute('SELECT Name ,Email , Section ,Salary , ID FROM Employees')
		self.data_sql = self.res.fetchall()

		for self.name , self.eamil , self.section , self.salary , self.id in self.data_sql:

			self.row = self.table.add_row().cells
			self.row[0].text = self.name
			self.row[1].text = self.eamil
			self.row[2].text = self.section
			self.row[3].text = self.salary
			self.row[4].text = str(self.id)


		
		self.path = filedialog.asksaveasfilename(initialdir = ".",title = "Save As",filetypes = ("docx files",'docx files(*.*)'))
		self.doc.save(self.path)
		self.con.commit()
		
