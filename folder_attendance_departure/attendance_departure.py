

from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from tkcalendar import Calendar, DateEntry
from PIL import Image , ImageTk
from datetime import datetime
import os , sys , shutil
import sqlite3
import time
from tkinter import filedialog
import csv
import docx
from threading import Thread

from babel.dates import format_date, parse_date, get_day_names, get_month_names
from babel.numbers import * 
from folder_attendance_departure.database_attendance_departure import DataBaseAttendanceDeparture

class AttendanceAndDeparture():
	def __init__(self,window):
		self.window = Toplevel()

		self.main_path  = os.path.dirname(os.path.abspath(__file__))
		self.path_database = os.path.join(self.main_path, "DataBaseAttendanceDeparture.db")
		self.window.iconbitmap('folder_attendance_departure\\images\\image_icon.ico')

		self.window.geometry("1180x800+300+100")
		self.window.resizable(width=False,height=False)
		self.window.configure(background='#EAF2F8')
		self.window.title("الحضور والإنصراف")



		self.menubar = Menu(self.window)
		self.file = Menu(self.menubar,tearoff=False)
		self.file.add_command(label='Save As')

		self.menubar.add_cascade(label='حفظ البيانات',command=self.save_list_employees)
		self.window.config(menu=self.menubar)




		self.class_database_attendance = DataBaseAttendanceDeparture(self.window)
		



		self.lable_title = Label(self.window,background='#E8E4FA',text='بيانات الحضور والإنصراف',width=71,font=("Libre Baskerville, serif;",20)).place(x=20,y=4)
		self.frame_info = Frame(self.window ,background='#D4E6F1',
		relief=RIDGE , bd=2 , width=1143,height=250).place(x=20,y=203)






		self.run_funcrion_add_frame = Thread(target=self.add_frame_search(),args=(),).start()
		self.search_employee = Thread(target=self.search_employee(),args=(),).start()


		self.add_frame_filters = Thread(target=self.add_frame_filters(),args=(),).start()
		self.filter_employee = Thread(target=self.filter_employee(),args=(),).start()
	
		self.add_id =  Thread(target=self.add_id(),args=(),).start()
		self.add_attendance_date_and_time_attendance = Thread(target=self.add_attendance_date_and_time_attendance(),args=(),).start() 



		self.function_add_leave_date = Thread(target=self.add_leave_date,args=())
		self.function_add_leave_date.start()

		self.function_add_notes = Thread(target=self.add_notes,args=())
		self.function_add_notes.start()


		self.function_add_name = Thread(target=self.add_name,args=())
		self.function_add_name.start()


		self.function_add_email = Thread(target=self.add_email,args=())
		self.function_add_email.start()


		self.function_add_frame_buttons = Thread(target=self.add_frame_buttons,args=())
		self.function_add_frame_buttons.start()
		

	


	def add_frame_search(self):
		""" The Function Add Entry Search Employee """

		self.frame_search = Frame(self.window,

			background='#d9ddf7', relief=RIDGE , bd=2 , width=1143,height=80
			).place(x=20,y=43)

		
		self.lable_search = Label(self.window,text="الرجاء إدخال رقم الموظف للبحث عنه",
			font=("bold",20),
			background='#d9ddf7' , 
			width=23
			).place(x=783,y=58)



		self.entry_search = Entry(
			self.window , width=45 ,font=('Bold Oblique',16) , relief=RAISED , bd=1 ,
			background='white',textvariable=self.class_database_attendance.variable_search_employee).place(x=225,y=62)



	def search_employee(self):
		self.image_search = Image.open('folder_attendance_departure\\images\\search.png')
		self.image_search = self.image_search.resize((35,35))
		self.insert_image_search = ImageTk.PhotoImage(self.image_search)

		self.button_search = Button(
		self.window , background='#F7F7F9' , width=90 , height=30 , text='بــحث',font=('Bold Oblique',18) ,
		image=self.insert_image_search , compound='left',padx=20,relief=FLAT , bd=1,
		command=self.class_database_attendance.search_employee).place(x=50,y=57)



	def add_frame_filters(self):
		self.frame_filters = Frame(self.window,

			background='#E5E1F4', relief=RIDGE , bd=2 , width=1143,height=80
			).place(x=20,y=123)

		
		self.lable_filter_one = Label(self.window,text="الرجاء إختيار التأريخ لفلترة البيانات من",
			font=("bold",20),
			background='#E5E1F4' , 
			width=23
			).place(x=783,y=139)


		self.filter_data_one = DateEntry(self.window,selectmode='day',year=datetime.now().year,
			month=datetime.now().month,day=datetime.now().day,
			font=('Bold Oblique',13),
			justify='center',date_pattern="yyyy-mm-dd", foreground="black",width=23,
			headersforeground="black", selectforeground="black",textvariable=self.class_database_attendance.filter_next).place(x=540,y=146)



		self.lable_filter_two = Label(self.window,text="إلى",
			font=("bold",20),
			background='#E5E1F4' , 
			width=5
			).place(x=440,y=139)

		self.filter_data_two = DateEntry(self.window,selectmode='day',year=datetime.now().year,
			month=datetime.now().month,day=datetime.now().day,
			font=('Bold Oblique',13),
			justify='center',date_pattern="yyyy-mm-dd", foreground="black",width=23,
			headersforeground="black", selectforeground="black",textvariable=self.class_database_attendance.filter_last).place(x=205,y=146)


	def filter_employee(self):
		self.image_filter = Image.open('folder_attendance_departure\\images\\filter_user.png')
		self.image_filter = self.image_filter.resize((35,35))
		self.insert_image_filter = ImageTk.PhotoImage(self.image_filter)

		self.button_search = Button(
		self.window , background='#F7F7F9' , width=90 , height=30 , text='فلترة',font=('Bold Oblique',18) ,
		image=self.insert_image_filter , compound='left',padx=20,relief=FLAT , bd=1,command=self.class_database_attendance.filters_data_employee).place(x=50,y=138)




		
		

	def add_email(self):
		label_email = Label(self.window,text='البريد الإلكتروني',font=("bold",20),bg='#D4E6F1').place(x=420,y=393)
		self.email = Entry(self.window , width=27,font=('Bold Oblique',18),
			relief=RIDGE , bd=1,justify='center',textvariable=self.class_database_attendance.var_email).place(x=50,y=395)



	def add_notes(self):
		self.label_notes = Label(self.window, text='ملاحضات',font=("bold",20),bg='#D4E6F1').place(x=420,y=333)

		self.notes = Entry(self.window , width=27,font=('Bold Oblique',18)
			,relief=RIDGE , bd=1,justify='center',
			textvariable=self.class_database_attendance.var_notes).place(x=50,y=335)





		# #-------------- Add Date Leave date -----------------
	def add_leave_date(self):
		self.label_leave_data = Label(self.window,text='تأريخ الإنصراف',font=("bold",20),bg='#D4E6F1').place(x=420,y=274)
		self.leave_data = DateEntry(self.window,selectmode='day',year=datetime.now().year,
				month=datetime.now().month,day=datetime.now().day,
				font=('Bold Oblique',16),
				justify='center',date_pattern="yyyy-mm-dd", foreground="black",width=27,
				headersforeground="black", selectforeground="black",textvariable=self.class_database_attendance.var_date_insraf).place(x=50,y=277)





		self.label_check_out_time = Label(self.window,text='وقت الإنصراف',font=("bold",20),bg='#D4E6F1').place(x=1010,y=333)
		self.check_out_time = Entry(self.window , width=28,font=('Bold Oblique',18),
			relief=RIDGE , bd=1,justify='center',textvariable=self.class_database_attendance.var_time_insraf).place(x=620,y=335)








	def add_attendance_date_and_time_attendance(self):
		self.label_attendance_date = Label(self.window,text='تأريخ الحضور',font=("bold",20),bg='#D4E6F1').place(x=420,y=216)
		self.attendance_date = DateEntry(self.window,selectmode='day',year=datetime.now().year,
				month=datetime.now().month,day=datetime.now().day,
				font=('Bold Oblique',16),
				justify='center',date_pattern="yyyy-mm-dd", foreground="black",width=27,
				headersforeground="black", selectforeground="black",textvariable=self.class_database_attendance.var_date_7dor).place(x=50,y=218)



		self.label_time_attendance = Label(self.window,text='وقت الحضور',font=("bold",20),bg='#D4E6F1').place(x=1020,y=274)
		self.time_attendance = Entry(self.window , width=30,font=('Bold Oblique',18)
			,relief=RIDGE , bd=1,justify='center',textvariable=self.class_database_attendance.var_time_7dor).place(x=620,y=276)






	def add_name(self):
		label_name = Label(self.window,text='إسم الموظف',font=("bold",20),bg='#D4E6F1').place(x=1012,y=395)
		self.name = Entry(self.window , width=27,font=('Bold Oblique',18),
			relief=RIDGE , bd=1,justify='center',textvariable=self.class_database_attendance.var_name).place(x=620,y=395)






	def add_id(self):
		self.label_id = Label(self.window,text='رقم الهوية',font=("bold",20),bg='#D4E6F1').place(x=1030,y=218)
		self.id = Entry(self.window , width=30,font=('Bold Oblique',18),
			relief=RIDGE , bd=1,justify='center',textvariable=self.class_database_attendance.var_id).place(x=620,y=220)







	def add_frame_buttons(self):
		self.frame_buttons = Frame(self.window,

			background='#E9E0F6', relief=RIDGE , bd=2 , width=1143,height=80
			).place(x=20,y=453)


		self.function_insert = Thread(target=self.function_insert , args=())
		self.function_insert.start()


		self.function_update = Thread(target=self.function_update , args=())
		self.function_update.start()


		self.function_delete = Thread(target=self.function_delete , args=())
		self.function_delete.start()


		self.function_view = Thread(target=self.function_view , args=())
		self.function_view.start()


		self.function_empty = Thread(target=self.function_empty , args=())
		self.function_empty.start()





	def function_insert(self):
		""" Insert Data In DataBase """
		self.image_insert = Image.open('folder_attendance_departure\\images\\add_user.png')
		self.image_insert = self.image_insert.resize((41,41))
		self.image_save = ImageTk.PhotoImage(self.image_insert)

		

		self.button_insert = Button(
			self.window , background='white' , width=145 , height=40 , text='حفظ البيانات',font=('Bold Oblique',18),
			relief=FLAT , bd=1 ,image=self.image_save , compound='left',padx=20,
			command=self.class_database_attendance.check_data_attendance).place(x=956,y=468)
 


	def function_update(self):
		""" Update Data In DataBase """
		self.image_update = Image.open('folder_attendance_departure\\images\\update_user.png')
		self.image_update = self.image_update.resize((41,41))
		self.image_save_update = ImageTk.PhotoImage(self.image_update)

		

		self.button_update = Button(
			self.window , background='white' , width=145 , height=40 , text='تحديث البيانات',font=('Bold Oblique',18),
			relief=FLAT , bd=1 ,image=self.image_save_update , compound='left',padx=20,
			command=self.class_database_attendance.update_data).place(x=730,y=468)
 




	def function_delete(self):
		""" Delete Data In DataBase """
		self.image_delete = Image.open('folder_attendance_departure\\images\\delete_user.png')
		self.image_delete = self.image_delete.resize((41,41))
		self.image_save_delete = ImageTk.PhotoImage(self.image_delete)

		

		self.button_delete = Button(
			self.window , background='white' , width=145 , height=40 , text='حذف البيانات',font=('Bold Oblique',18),
			relief=FLAT , bd=1 ,image=self.image_save_delete , compound='left',padx=20,
			command=self.class_database_attendance.delete_data).place(x=500,y=468)
 	




	def function_view(self):
		""" view Data In DataBase """
		self.image_view = Image.open('folder_attendance_departure\\images\\view_data.png')
		self.image_view = self.image_view.resize((39,39))
		self.image_save_view = ImageTk.PhotoImage(self.image_view)

		

		self.button_view = Button(
			self.window , background='white' , width=145 , height=40 , text='عرض البيانات',font=('Bold Oblique',18),
			relief=FLAT , bd=1 ,image=self.image_save_view , compound='left',padx=20,
			command=self.class_database_attendance.show_data).place(x=280,y=468)
 


	def function_empty(self):
		""" view Data In DataBase """
		self.image_empty = Image.open('folder_attendance_departure\\images\\empty_data.png')
		self.image_empty = self.image_empty.resize((39,39))
		self.image_save_empty = ImageTk.PhotoImage(self.image_empty)

		

		self.button_empty = Button(
			self.window , background='white' , width=145 , height=40 , text='إفراغ الحقول',font=('Bold Oblique',18),
			relief=FLAT , bd=1 ,image=self.image_save_empty , compound='left',padx=20,
			command=self.class_database_attendance.clean_data).place(x=50,y=468)
 





	def save_list_employees(self):
		self.title = "قائمة الحضور والإنصراف بتأريخ"
		self.date_time = str(datetime.now().date())


		self.doc = docx.Document()
		self.doc.add_heading(f"{self.title} - {self.date_time}")
		

		self.table = self.doc.add_table(rows=1, cols=9)
		self.row = self.table.rows[0].cells
		self.table.style = "Colorful List Accent 1"

		self.table_header = [
							"البريد الإلكتلاوني" , "ملاحضات" , "تاريخ الإنصراف",
							"تاريخ الحضور" , "الإسم" , "وقت الإنصراف",
							"وقت الحضور" , "الرقم الوظيفي"

							]



		for i in range(8):
			self.table.rows[0].cells[i].text = self.table_header[i]


		self.con = sqlite3.connect(self.path_database)
		self.curs = self.con.cursor()

		self.res = self.curs.execute("SELECT Email , Notes , date_ansraf , date_7dor ,Name , time_ansraf,time_7dor,ID FROM AttendanceDeparture")
		self.data_sql = self.res.fetchall()

		for self.email , self.notes , self.date_insraf  , self.date_7dor , self.name ,self.time_ansraf , self.time_7dor , self.id_ in self.data_sql:

			self.row = self.table.add_row().cells
			self.row[0].text = self.email
			self.row[1].text = self.notes
			self.row[2].text = str(self.date_insraf)
			self.row[3].text = str(self.date_7dor)
			self.row[4].text = str(self.name)
			self.row[5].text = str(self.time_ansraf)

			self.row[6].text = str(self.time_7dor)
			self.row[7].text = str(self.id_)
		

		try :
		
			self.path = filedialog.asksaveasfilename(initialdir = ".",title = "Save As",filetypes = [("Docx files", ".docx .docx")])
			self.doc.save(self.path)
			self.con.commit()
		except :
			pass
